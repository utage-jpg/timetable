# build_schedule.py
# 時間表 .docx ファイルを解析して schedule_data.js を生成します
# 出力先: ../schedule_data.js

import sys, os, re, json
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("python-docx が必要です: pip install python-docx")
    sys.exit(1)

SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, '..', 'schedule_data.js')

# 月 → 年
def month_to_year(m):
    return 2025 if m >= 11 else 2026

# ファイル名 → コース名
def course_name_from_path(path):
    base = os.path.splitext(os.path.basename(path))[0]
    base = re.sub(r'^【[^】]*】', '', base).strip()
    # "2026 ", "2026-", "2026年" などの年プレフィックスを除去
    base = re.sub(r'^2026[-年\s　]*', '', base).strip()
    base = re.sub(r'\s*\(\d+\)\s*$', '', base).strip()
    return base

# ----------------------------------------------------------------
# ヘッダー正規化 (空白・全角空白を除去)
# ----------------------------------------------------------------
def norm_header(s):
    return re.sub(r'[\s\u3000\xa0]+', '', s)

# ----------------------------------------------------------------
# 列インデックス検出
# ----------------------------------------------------------------
def detect_columns(header_row):
    norm = [norm_header(c) for c in header_row]

    # 日程列: "日程", "日 程", "講義日" などを検出
    date_col = None
    for i, v in enumerate(norm):
        if '日程' in v or v == '講義日':
            date_col = i
            break

    # 時間列: "講義時間" を優先
    time_col = None
    for i, v in enumerate(norm):
        if '講義時間' in v:
            time_col = i
            break
    if time_col is None:
        for i, v in enumerate(norm):
            if '時間' in v:
                time_col = i
                break

    # 科目列: 優先度: 科目 > 内容 > 講義(単独)
    subject_col = None
    for i, v in enumerate(norm):
        if '科目' in v:
            subject_col = i
            break
    if subject_col is None:
        for i, v in enumerate(norm):
            if '内容' in v:
                subject_col = i
                break
    if subject_col is None:
        # "講義" のみ（"講義時間" "講義日" を除く）
        for i, v in enumerate(norm):
            if v == '講義':
                subject_col = i
                break

    return date_col, time_col, subject_col

# ----------------------------------------------------------------
# 日付パーサー
# ----------------------------------------------------------------
# フル日付: "11／22 (土)", "4／ 26 (日)", "3／01（日）", "1/12(月)"
_FULL_DATE_RE = re.compile(
    r'(\d{1,2})[\/／]\s*(\d{1,2})\s*[\(（]?[月火水木金土日][\)）]?'
)
# 日のみ: "13 (火)", "9(火)", "13（火）"
_DAY_ONLY_RE = re.compile(
    r'^(\d{1,2})\s*[\(（][月火水木金土日][\)）]$'
)

def parse_date_cell(cell_text, carry_month):
    t = cell_text.strip().replace('\u3000', ' ').replace('\xa0', ' ')

    if not t or t.lower().startswith('web'):
        return None, carry_month

    m = _FULL_DATE_RE.match(t)
    if m:
        month = int(m.group(1))
        day   = int(m.group(2))
        year  = month_to_year(month)
        return f'{year}-{month:02d}-{day:02d}', month

    m2 = _DAY_ONLY_RE.match(t)
    if m2 and carry_month is not None:
        day  = int(m2.group(1))
        year = month_to_year(carry_month)
        return f'{year}-{carry_month:02d}-{day:02d}', carry_month

    return None, carry_month

# ----------------------------------------------------------------
# 時間パーサー
# ----------------------------------------------------------------
def parse_time_cell(t):
    t = t.strip().replace('：', ':').replace('〜', '~').replace('～', '~')
    m = re.search(r'(\d{1,2}:\d{2})\s*[~～〜]\s*(\d{1,2}:\d{2})', t)
    if m:
        return m.group(1) + '〜' + m.group(2)
    return t or None

# ----------------------------------------------------------------
# セクションヘッダー判定（全セルが同じ非空値）
# ----------------------------------------------------------------
def is_section_header(cells):
    vals = [c.strip() for c in cells if c.strip()]
    return len(vals) > 0 and len(set(vals)) == 1

# ----------------------------------------------------------------
# 科目名クリーニング
# ----------------------------------------------------------------
def clean_subject(t):
    t = t.strip()
    # 先頭の "(完)" "(養)" "●" などを除去
    t = re.sub(r'^[●▲▼◆◇■□★☆※]', '', t)
    t = re.sub(r'^[\(（][^）\)]{1,4}[\)）]\s*', '', t)
    t = re.sub(r'[\u3000\xa0　]+', ' ', t)
    return t.strip() or None

# ----------------------------------------------------------------
# 設計製図パーフェクト本科: 特殊形式の解析
# 列構成: [セクション名, 講義日, 回数, 回数, 講義時間, 講義, 講義, 宿題]
# 行ごとにセクション名が入っており、セクションヘッダー行は存在しない
# ----------------------------------------------------------------
def parse_sekkei_docx(path, course):
    doc    = Document(path)
    events = []
    if not doc.tables:
        return events

    tbl  = doc.tables[0]
    rows = tbl.rows
    if not rows:
        return events

    # ヘッダー行を確認
    header = [norm_header(c.text) for c in rows[0].cells]
    # 期待: [空, 講義日, 回数, 回数, 講義時間, 講義, 講義, ...]
    # date=1, time=4, subject=5

    date_col    = next((i for i,v in enumerate(header) if v == '講義日'), None)
    time_col    = next((i for i,v in enumerate(header) if '講義時間' in v), None)
    subject_col = None
    for i, v in enumerate(header):
        if v == '講義' and i != time_col:
            subject_col = i
            break

    if date_col is None or subject_col is None:
        print(f"  [WARN] 設計製図列検出失敗: {os.path.basename(path)}")
        return events

    carry_month = None
    for row in rows[1:]:
        cells = [c.text for c in row.cells]
        if len(cells) <= subject_col:
            continue

        date_text    = cells[date_col].strip()
        time_text    = cells[time_col].strip() if time_col is not None else ''
        subject_text = cells[subject_col].strip()

        if not date_text or date_text.lower().startswith('web'):
            continue

        date_iso, carry_month = parse_date_cell(date_text, carry_month)
        if date_iso is None:
            continue

        subject_str = clean_subject(subject_text)
        if not subject_str:
            continue

        events.append({
            'date':    date_iso,
            'time':    parse_time_cell(time_text) if time_text else None,
            'subject': subject_str,
            'course':  course,
        })

    return events

# ----------------------------------------------------------------
# 1ファイル解析
# ----------------------------------------------------------------
def parse_docx(path):
    course = course_name_from_path(path)
    doc    = Document(path)
    events = []

    if not doc.tables:
        return course, events

    tbl     = doc.tables[0]
    rows    = tbl.rows
    if not rows:
        return course, events

    header  = [c.text for c in rows[0].cells]
    date_col, time_col, subject_col = detect_columns(header)

    # 設計製図は専用パーサー
    if date_col is None and any('講義日' in norm_header(h) for h in header):
        return course, parse_sekkei_docx(path, course)

    if date_col is None or subject_col is None:
        print(f"  [WARN] 列検出失敗: {os.path.basename(path)} header={[norm_header(h) for h in header]}")
        return course, events

    carry_month = None

    for row in rows[1:]:
        cells = [c.text for c in row.cells]
        if len(cells) <= subject_col:
            continue

        if is_section_header(cells):
            continue

        date_text    = cells[date_col].strip() if date_col < len(cells) else ''
        time_text    = cells[time_col].strip() if time_col is not None and time_col < len(cells) else ''
        subject_text = cells[subject_col].strip()

        if not date_text or date_text.lower().startswith('web'):
            continue

        date_iso, carry_month = parse_date_cell(date_text, carry_month)
        if date_iso is None:
            continue

        subject_str = clean_subject(subject_text)
        if not subject_str:
            continue

        events.append({
            'date':    date_iso,
            'time':    parse_time_cell(time_text) if time_text else None,
            'subject': subject_str,
            'course':  course,
        })

    return course, events

# ----------------------------------------------------------------
# メイン
# ----------------------------------------------------------------
def main():
    docx_files = sorted([
        os.path.join(SCRIPT_DIR, f)
        for f in os.listdir(SCRIPT_DIR)
        if f.endswith('.docx')
    ])

    all_events  = []
    all_courses = []

    for path in docx_files:
        fname = os.path.basename(path)
        print(f"解析中: {fname}")
        try:
            course, events = parse_docx(path)
            print(f"  → {course}: {len(events)} 件")
            all_events.extend(events)
            if course not in all_courses:
                all_courses.append(course)
        except Exception as e:
            import traceback
            print(f"  [ERROR] {fname}: {e}")
            traceback.print_exc()

    all_events.sort(key=lambda e: e['date'])

    js = (
        '// schedule_data.js — 自動生成ファイル (build_schedule.py)\n'
        '// 手動編集しないでください\n\n'
        'const SCHEDULE_DATA = '
        + json.dumps({'courses': all_courses, 'events': all_events},
                     ensure_ascii=False, indent=2)
        + ';\n'
    )

    out = os.path.normpath(OUTPUT_PATH)
    with open(out, 'w', encoding='utf-8') as f:
        f.write(js)

    print(f'\n完了: {len(all_events)} 件 → {out}')

if __name__ == '__main__':
    main()
